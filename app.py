import os
import tempfile
from docx import Document
from google import genai
import streamlit as st

st.set_page_config(
    page_title="Word to LaTeX Converter", page_icon="📄", layout="centered"
)

st.title("📄 Word Document to LaTeX Converter")
st.write(
    "Upload a Word document (.docx) to convert its structure cleanly into"
    " LaTeX."
)

api_key = st.secrets.get("GEMINI_API_KEY") or os.environ.get("GEMINI_API_KEY")

if not api_key:
  st.error("Gemini API key not found. Please configure it in Streamlit Secrets.")
else:
  client = genai.Client(api_key=api_key)

  uploaded_file = st.file_uploader(
      "Choose a Word document", type=["docx", "txt"]
  )

  if uploaded_file is not None:
    with tempfile.NamedTemporaryFile(
        delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}"
    ) as tmp_file:
      tmp_file.write(uploaded_file.getvalue())
      tmp_path = tmp_file.name

    if st.button("Convert to LaTeX"):
      with st.spinner("Extracting structure and generating LaTeX..."):
        try:
          document_content = ""
          if uploaded_file.name.endswith(".docx"):
            doc = Document(tmp_path)
            # Extract paragraphs and tables sequentially
            fullText = []
            for para in doc.paragraphs:
              if para.text.strip():
                fullText.append(para.text)
            for table in doc.tables:
              for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                fullText.append(" | ".join(row_text))
            document_content = "\n".join(fullText)
          else:
            with open(tmp_path, "r", encoding="utf-8") as f:
              document_content = f.read()

          prompt = f"""
                    You are an expert LaTeX typesetter. 
                    Convert the following extracted document text and structure into clean, complete, and compilable LaTeX code.
                    Use an appropriate document class (like article), proper sectioning commands (\\section, \\subsection), 
                    and format any lists or tables correctly.
                    Return ONLY the raw LaTeX code inside a markdown code block.

                    Document Content:
                    {document_content}
                    """

          response = client.models.generate_content(
              model="gemini-3.6-flash", contents=prompt
          )

          latex_output = (
              response.text.replace("```latex", "")
              .replace("```", "")
              .strip()
          )

          st.success("Conversion successful!")
          st.code(latex_output, language="latex")

          st.download_button(
              label="Download .tex File",
              data=latex_output,
              file_name="output.tex",
              mime="text/plain",
          )

        except Exception as e:
          st.error(f"An error occurred: {e}")
        finally:
          if os.path.exists(tmp_path):
            os.remove(tmp_path)
